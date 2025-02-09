import frappe
import csv, os, tempfile, json
import openpyxl
from docx import Document
from reportlab.pdfgen import canvas
from frappe.utils.file_manager import save_file
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle

def _export_data(fields, only_active, file_format="csv", job_id=None):
    """
    Exportiert Mitgliederliste in verschiedenen Formaten.
    
    Unterstützte Formate: CSV, XLSX, DOCX, PDF, TXT
    """
    fields = json.loads(fields) if isinstance(fields, str) else fields
    only_active = only_active.lower() == "true" if isinstance(only_active, str) else only_active

    filters = {"status": "Aktiv"} if only_active else {}
    mitglieder = frappe.get_all("Mitglied", filters=filters, fields=fields)

    if not mitglieder:
        frappe.throw("Es wurden keine Mitglieder gefunden, die exportiert werden können.")

    temp_dir = tempfile.gettempdir()
    filename = f"Mitgliederliste_{frappe.session.user}_{frappe.utils.now_datetime().strftime('%Y%m%d%H%M%S')}.{file_format}"
    file_path = os.path.join(temp_dir, filename)

    frappe.logger().info(f"📂 Speichere Datei unter: {file_path}")

    try:
        if file_format == "csv":
            _export_csv(file_path, fields, mitglieder)
        elif file_format in ["xls", "xlsx"]:
            _export_xlsx(file_path, fields, mitglieder)
        elif file_format in ["doc", "docx"]:
            _export_docx(file_path, fields, mitglieder)
        elif file_format == "pdf":
            _export_pdf(file_path, fields, mitglieder)
        elif file_format == "txt":
            _export_txt(file_path, fields, mitglieder)
        else:
            frappe.throw("Ungültiges Exportformat.")

        with open(file_path, mode="rb") as file:
            file_content = file.read()

        file_doc = save_file(filename, file_content, "Mitglied", "Mitgliederliste-Export", is_private=1)
        os.remove(file_path)

        # ✅ Echtzeit-Event senden
        frappe.publish_realtime("export_complete", {"status": "success", "file_url": file_doc.file_url})

        # ✅ Status für Fortschrittsanzeige setzen
        if job_id:
            frappe.cache().set_value(f"export_status_{job_id}", "success")

        return file_doc.file_url

    except Exception as e:
        frappe.logger().error(f"❌ Fehler beim Export: {str(e)}")

        if job_id:
            frappe.cache().set_value(f"export_status_{job_id}", "error")

        frappe.throw(f"Fehler beim Export: {str(e)}")

# 🔧 **1. CSV-Export**
def _export_csv(file_path, fields, mitglieder):
    with open(file_path, mode="w", newline="", encoding="utf-8") as file:
        writer = csv.writer(file)
        writer.writerow(fields)
        for mitglied in mitglieder:
            writer.writerow([mitglied[field] for field in fields])

# 🔧 **2. XLSX-Export**
def _export_xlsx(file_path, fields, mitglieder):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(fields)
    for mitglied in mitglieder:
        ws.append([mitglied[field] for field in fields])
    wb.save(file_path)

# 🔧 **3. DOCX-Export**
def _export_docx(file_path, fields, mitglieder):
    doc = Document()
    doc.add_heading("Mitgliederliste", level=1)
    table = doc.add_table(rows=1, cols=len(fields))
    hdr_cells = table.rows[0].cells
    for i, field in enumerate(fields):
        hdr_cells[i].text = field
    for mitglied in mitglieder:
        row_cells = table.add_row().cells
        for i, field in enumerate(fields):
            row_cells[i].text = str(mitglied[field])
    doc.save(file_path)

# 🔧 **4. PDF-Export**
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle

def _export_pdf(file_path, fields, mitglieder):
    """
    Erstellt eine formatierte PDF-Tabelle für den Mitgliederexport.
    - Passt sich automatisch an die Seitenbreite an.
    - Setzt das Format auf Querformat, falls nötig.
    """
    # 🔹 Querformat für mehr Platz
    doc = SimpleDocTemplate(file_path, pagesize=landscape(A4))
    elements = []

    # 🔹 Kopfzeile (Überschriften)
    data = [fields]  # Erste Zeile = Spaltenüberschriften

    # 🔹 Daten einfügen
    for mitglied in mitglieder:
        data.append([str(mitglied.get(field, "")) for field in fields])

    # 🔹 Tabellenbreite berechnen (max. Breite = Seitenbreite - Rand)
    page_width = landscape(A4)[0]  # Breite der Seite (Querformat)
    left_margin = right_margin = 40  # Ränder für den Druck
    max_table_width = page_width - left_margin - right_margin

    # 🔹 Dynamische Spaltenbreiten berechnen
    num_cols = len(fields)
    col_width = max_table_width / num_cols
    col_widths = [col_width] * num_cols  # Alle Spalten gleichmäßig aufteilen

    # 🔹 Tabelle erstellen
    table = Table(data, colWidths=col_widths)

    # 🔹 Tabellenstil definieren
    style = TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),  # Kopfzeile grau hinterlegen
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),  # Schrift in Kopfzeile weiß
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),  # Links ausgerichtet
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),  # Kopfzeile fett
        ('BOTTOMPADDING', (0, 0), (-1, 0), 10),  # Mehr Abstand zur ersten Zeile
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),  # Restliche Zeilen mit Hintergrundfarbe
        ('GRID', (0, 0), (-1, -1), 1, colors.black)  # Rahmen um jede Zelle
    ])
    table.setStyle(style)

    elements.append(table)
    doc.build(elements)

# 🔧 **5. TXT-Export**
def _export_txt(file_path, fields, mitglieder):
    with open(file_path, mode="w", encoding="utf-8") as file:
        file.write("\t".join(fields) + "\n")
        for mitglied in mitglieder:
            file.write("\t".join(str(mitglied[field]) for field in fields)) + "\n"


@frappe.whitelist()
def export_data_async(fields="[]", only_active="false", file_format="csv"):
    """Startet den Export mit Status-Tracking"""

    job_id = frappe.generate_hash(length=10)
    frappe.cache().set_value(f"export_status_{job_id}", "started")

    try:
        job = frappe.enqueue(
            "spherdex.global_scripts.export_utils._export_data",
            queue="long",
            timeout=600,
            job_name=f"Mitglieder-{file_format}-Export",
            is_async=True,
            fields=fields,
            only_active=only_active,
            file_format=file_format,
            job_id=job_id
        )
        return {"status": "Export gestartet", "job_id": job_id}

    except Exception as e:
        frappe.cache().set_value(f"export_status_{job_id}", "error")
        return {"status": "Fehler", "message": str(e)}

@frappe.whitelist()
def is_export_ready():
    """
    Prüft, ob der Export abgeschlossen ist.

    Returns:
        dict: {"export_ready": True/False}
    """
    ready = frappe.cache().get_value("export_ready") or False
    return {"export_ready": ready}

@frappe.whitelist()
def delete_export_files():
    """
    Löscht alle gespeicherten Exportdateien und die zugehörigen Datenbankeinträge.
    
    ERPNext speichert pro Datei zwei Einträge:
    1. Die eigentliche Datei im `private/files/`-Ordner
    2. Einen Eintrag in `tabFile` für die Datei-Verwaltung.
    
    Da wir eindeutige Namen nutzen, löschen wir alle Dateien mit dem gleichen Präfix.
    """
    base_filename = "Mitgliederliste_"
    site_path = frappe.get_site_path("private/files/")

    try:
        # 🔍 Alle passenden Dateien im Verzeichnis finden
        matching_files = [f for f in os.listdir(site_path) if f.startswith(base_filename)]

        if not matching_files:
            frappe.logger().warning("⚠️ Keine Exportdateien gefunden.")
            return {"status": "success", "message": "Keine Dateien gefunden."}

        # 🔥 Alle gefundenen Dateien löschen
        for file in matching_files:
            file_path = os.path.join(site_path, file)
            if os.path.exists(file_path):
                os.remove(file_path)
                frappe.logger().info(f"🗑 Datei gelöscht: {file_path}")

        # 🔥 Dazugehörige Einträge aus `tabFile` löschen
        file_records = frappe.get_all("File", filters={"file_name": ("like", base_filename + "%")}, fields=["name"])
        
        if file_records:
            for file_record in file_records:
                frappe.delete_doc("File", file_record["name"], force=1)
            frappe.db.commit()
            frappe.logger().info(f"🗑 {len(file_records)} Datenbankeintrag(e) in `tabFile` gelöscht.")

        return {"status": "success", "message": f"{len(matching_files)} Datei(en) und {len(file_records)} Datenbankeintrag(e) gelöscht."}

    except Exception as e:
        frappe.logger().error(f"❌ Fehler beim Löschen: {str(e)}")
        return {"status": "error", "message": str(e)}


@frappe.whitelist()
def get_export_status(job_id):
    """Gibt den aktuellen Status des Exports zurück"""
    status = frappe.cache().get_value(f"export_status_{job_id}") or "unknown"
    return {"job_id": job_id, "status": status}