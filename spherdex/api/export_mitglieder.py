import frappe
import csv, os, tempfile, json
import openpyxl
from docx import Document
from reportlab.pdfgen import canvas
from frappe.utils.file_manager import save_file
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.utils import simpleSplit
from reportlab.lib.units import cm

@frappe.whitelist()
def _export_data(fields, only_active, file_format="csv", job_id=None, member_id=None):
    """Exportiert Mitgliederliste oder einzelne Datensätze mit allen Spalten."""
    
    is_single_member = bool(member_id)  # 🔍 Prüft, ob `member_id` übergeben wurde

    # ✅ **Daten abrufen: Entweder alle Mitglieder oder nur ein einzelnes**
    if is_single_member:
        mitglieder = [frappe.get_doc("Mitglied", member_id).as_dict()]
    else:
        filters = {"status": "Aktiv"} if only_active else {}
        mitglieder = frappe.get_all("Mitglied", filters=filters, fields=fields)

    if not mitglieder:
        frappe.throw("Es wurden keine Mitglieder gefunden, die exportiert werden können.")

    temp_dir = tempfile.gettempdir()
    filename = (
        f"mitglied_{member_id}.{file_format}"
        if is_single_member else
        f"Mitgliederliste_{frappe.session.user}_{frappe.utils.now_datetime().strftime('%Y%m%d%H%M%S')}.{file_format}"
    )
    file_path = os.path.join(temp_dir, filename)

    # ✅ **Hier `data` korrekt definieren**
    data = [fields]  # Erste Zeile: Spaltenüberschriften
    for mitglied in mitglieder:
        row = [str(mitglied.get(field, "")) for field in fields]
        data.append(row)

    try:
        # 🔹 **Richtige Export-Funktion basierend auf `is_single_member`**
        if file_format == "pdf":
            _export_pdf(file_path, fields, mitglieder, is_single_member)
        elif file_format == "csv":
            _export_csv(file_path, fields, mitglieder)
        elif file_format in ["xls", "xlsx"]:
            _export_xlsx(file_path, fields, mitglieder)
        elif file_format in ["doc", "docx"]:
            _export_docx(file_path, fields, mitglieder)
        else:
            frappe.throw("Ungültiges Exportformat.")

        with open(file_path, mode="rb") as file:
            file_content = file.read()

        file_doc = save_file(filename, file_content, "Mitglied", "Mitgliederliste-Export", is_private=1)
        os.remove(file_path)

        frappe.publish_realtime("export_complete", {"status": "success", "file_url": file_doc.file_url})

        if job_id:
            frappe.cache().set_value(f"export_status_{job_id}", "success")

        return file_doc.file_url

    except Exception as e:
        frappe.logger().error(f"❌ Fehler beim Export: {str(e)}")

        if job_id:
            frappe.cache().set_value(f"export_status_{job_id}", "error")

        frappe.throw(f"Fehler beim Export: {str(e)}")

# 🔧 **1. CSV-Export**
def _export_csv(file_path, fields, mitglieder):
    with open(file_path, mode="w", newline="", encoding="utf-8") as file:
        writer = csv.writer(file)
        writer.writerow(fields)
        for mitglied in mitglieder:
            writer.writerow([mitglied[field] for field in fields])

# 🔧 **2. XLSX-Export**
def _export_xlsx(file_path, fields, mitglieder):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(fields)
    for mitglied in mitglieder:
        ws.append([mitglied[field] for field in fields])
    wb.save(file_path)

# 🔧 **3. DOCX-Export**
def _export_docx(file_path, fields, mitglieder):
    doc = Document()
    doc.add_heading("Mitgliederliste", level=1)
    table = doc.add_table(rows=1, cols=len(fields))
    hdr_cells = table.rows[0].cells
    for i, field in enumerate(fields):
        hdr_cells[i].text = field
    for mitglied in mitglieder:
        row_cells = table.add_row().cells
        for i, field in enumerate(fields):
            row_cells[i].text = str(mitglied[field])
    doc.save(file_path)

# 🔧 **4. PDF-Export**
def _export_pdf(file_path, fields, mitglieder, is_single_member=False):
    """ Erstellt eine formatierte PDF mit vertikaler (Einzelmitglied) oder horizontaler (Liste) Tabelle. """

    # 🔹 Layout je nach Exporttyp (Einzeln = Hochformat / Liste = Querformat)
    page_size = A4 if is_single_member else landscape(A4)
    doc = SimpleDocTemplate(file_path, pagesize=page_size)
    elements = []

    if is_single_member:
        # 🔹 **Einzelmitglied: Vertikale Tabelle mit Spaltennamen in der linken Spalte**
        data = [[field, mitglieder[0].get(field, "—")] for field in fields]
        col_widths = [6 * cm, 10 * cm]  # Breite für "Feldname" & "Wert"
    else:
        # 🔹 **Liste: Horizontale Tabelle mit Spaltennamen als Kopfzeile**
        data = [fields]  # Kopfzeile
        for mitglied in mitglieder:
            data.append([mitglied.get(field, "—") for field in fields])
        col_widths = None  # Automatische Breite

    # 🔹 Tabelle erstellen
    table = Table(data, colWidths=col_widths)

    # 🔹 Stil für bessere Lesbarkeit
    style = TableStyle([
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke if not is_single_member else colors.black),  # Kopfzeile / Feldnamen weiß
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey if not is_single_member else None),  # Grauer Hintergrund für Kopfzeile (Liste)
        ('BACKGROUND', (0, 0), (0, -1), colors.grey if is_single_member else None),  # Grauer Hintergrund für Feldnamen (Einzelexport)
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),  # Fett für Kopfzeile (Liste)
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold' if is_single_member else 'Helvetica'),  # Fett für Feldnamen (Einzelexport)
        ('TEXTCOLOR', (0, 0), (0, -1), colors.whitesmoke if is_single_member else colors.black),  # Weiße Schrift für Feldnamen (Einzelexport)
        ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ])

    table.setStyle(style)

    # 🔹 Tabelle ins PDF einfügen
    elements.append(table)
    doc.build(elements)

# 🔧 **5. TXT-Export**
def _export_txt(file_path, fields, mitglieder):
    with open(file_path, mode="w", encoding="utf-8") as file:
        file.write("\t".join(fields) + "\n")
        for mitglied in mitglieder:
            file.write("\t".join(str(mitglied[field]) for field in fields)) + "\n"


@frappe.whitelist()
def export_data_async(fields="[]", only_active="false", file_format="csv", member_id=None):
    """Startet den Export für ein einzelnes Mitglied oder alle Mitglieder mit Status-Tracking."""

    job_id = frappe.generate_hash(length=10)
    frappe.cache().set_value(f"export_status_{job_id}", "started")

    # 🛠 **Feld-Liste sicher umwandeln**
    fields = json.loads(fields) if isinstance(fields, str) else fields

    # 🛠 **Falls `member_id` gesetzt ist, alle Felder automatisch abrufen**
    if member_id:
        meta = frappe.get_meta("Mitglied")
        fields = [df.fieldname for df in meta.fields if df.fieldtype not in ("Section Break", "Column Break", "HTML", "Table", "Button")]

    if not fields or len(fields) == 0:
        frappe.throw("⚠ Kein Feld zum Exportieren verfügbar.")

    try:
        frappe.enqueue("spherdex.api.export_mitglieder._export_data",
            queue="long",
            timeout=600,
            job_name=f"Mitglieder-{file_format}-Export",
            fields=fields,  # ✅ Direkt übergeben
            only_active=only_active,
            file_format=file_format,
            member_id=member_id,
            job_id=job_id
        )

        return {"status": "Export gestartet", "job_id": job_id}

    except Exception as e:
        frappe.cache().set_value(f"export_status_{job_id}", "error")
        return {"status": "Fehler", "message": str(e)}


@frappe.whitelist()
def is_export_ready():
    """
    Prüft, ob der Export abgeschlossen ist.

    Returns:
        dict: {"export_ready": True/False}
    """
    ready = frappe.cache().get_value("export_ready") or False
    return {"export_ready": ready}

@frappe.whitelist()
def delete_export_files():
    """
    Löscht gespeicherte Exportdateien erst nach erfolgreichem Download.
    """
    base_filenames = ["Mitgliederliste_", "mitglied_"]
    site_path = frappe.get_site_path("private/files/")

    try:
        # 🔍 Alle passenden Dateien im Verzeichnis finden
        matching_files = [
            f for f in os.listdir(site_path) if any(f.startswith(prefix) for prefix in base_filenames)
        ]

        if not matching_files:
            frappe.logger().warning("⚠️ Keine Exportdateien gefunden.")
            return {"status": "success", "message": "Keine Dateien gefunden."}

        # 🔥 Alle gefundenen Dateien löschen
        for file in matching_files:
            file_path = os.path.join(site_path, file)
            if os.path.exists(file_path):
                os.remove(file_path)
                frappe.logger().info(f"🗑 Datei gelöscht: {file_path}")

        # 🔥 Dazugehörige Einträge aus `tabFile` löschen
        file_records = frappe.get_all(
            "File", filters={"file_name": ("like", "Mitgliederliste_%")}, fields=["name"]
        ) + frappe.get_all(
            "File", filters={"file_name": ("like", "mitglied_%")}, fields=["name"]
        )

        if file_records:
            for file_record in file_records:
                frappe.delete_doc("File", file_record["name"], force=1)
            frappe.db.commit()
            frappe.logger().info(f"🗑 {len(file_records)} Datenbankeintrag(e) in `tabFile` gelöscht.")

        return {"status": "success", "message": f"{len(matching_files)} Datei(en) und {len(file_records)} Datenbankeintrag(e) gelöscht."}

    except Exception as e:
        frappe.logger().error(f"❌ Fehler beim Löschen: {str(e)}")
        return {"status": "error", "message": str(e)}


@frappe.whitelist()
def get_export_status(job_id):
    """Gibt den aktuellen Status des Exports zurück"""
    status = frappe.cache().get_value(f"export_status_{job_id}") or "unknown"
    return {"job_id": job_id, "status": status}

def process_member_export(member_id, file_format, job_id):
    """ Erstellt die Exportdatei und sendet eine Benachrichtigung, wenn sie fertig ist. """
    member = frappe.get_doc("Mitglied", member_id)

    # 📝 Alle Felder des Mitglieds-Typs abrufen
    data = [["Feldname", "Wert"]]
    for fieldname, value in member.as_dict().items():
        data.append([fieldname, str(value) if value else ""])

    # 🔗 Dynamisch alle verknüpften Tabellen abrufen
    child_tables = frappe.get_all(
        "DocField",
        filters={"fieldtype": "Table", "parent": "Mitglied"},
        fields=["options"]
    )

    for table in child_tables:
        doctype = table["options"]
        if frappe.db.exists("DocType", doctype):
            entries = frappe.get_all(doctype, filters={"parent": member_id}, fields=["*"])
            if entries:
                data.append(["---", "---"])
                data.append([f"📂 {doctype}-Einträge", ""])
                for entry in entries:
                    for key, value in entry.items():
                        data.append([key, str(value) if value else ""])

    # 📂 Datei im `private/files`-Ordner speichern
    filename = f"mitglied_{member.name}.{file_format}"
    file_path = frappe.get_site_path("private", "files", filename)

    # 📌 Richtige Export-Funktion aufrufen
    if file_format == "csv":
        _export_csv(file_path, [row[0] for row in data], [{row[0]: row[1] for row in data}])
    elif file_format == "pdf":
        _export_pdf(file_path, [row[0] for row in data], [{row[0]: row[1] for row in data}])
    elif file_format == "docx":
        _export_docx(file_path, [row[0] for row in data], [{row[0]: row[1] for row in data}])
    elif file_format == "xlsx":
        _export_xlsx(file_path, [row[0] for row in data], [{row[0]: row[1] for row in data}])
    elif file_format == "txt":
        _export_txt(file_path, [row[0] for row in data], [{row[0]: row[1] for row in data}])
    else:
        frappe.throw("Ungültiges Exportformat!")

    # ✅ Status für Fortschrittsanzeige setzen
    frappe.cache().set_value(f"export_status_{job_id}", "success")